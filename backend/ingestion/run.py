from __future__ import annotations

import argparse
import hashlib
import json
import os
import shutil
import subprocess
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pypdfium2 as pdfium
import pytesseract
from docx import Document
from PIL import Image
from pypdf import PdfReader
from tqdm import tqdm

try:
    import fitz  # PyMuPDF
except ImportError:  # pragma: no cover - optional fallback renderer
    fitz = None

try:
    from docling.document_converter import DocumentConverter
except ImportError:  # pragma: no cover - optional parser
    DocumentConverter = None


SUPPORTED_EXTENSIONS = {".pdf"}
WINDOWS_TESSERACT_DEFAULT = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")
DEFAULT_WORKERS = max(1, min(8, os.cpu_count() or 4))


@dataclass
class IngestionArgs:
    input_dir: Path
    output_dir: Path
    min_native_chars: int
    ocr_dpi: int
    ocr_lang: str
    workers: int
    tessdata_dir: Path | None
    disable_ocr: bool
    limit_docs: int | None
    pdf_parser: str


def parse_args() -> IngestionArgs:
    parser = argparse.ArgumentParser(
        description="Step 1-4 legal ingestion pipeline: checks, inventory, extraction, OCR fallback."
    )
    parser.add_argument("--input-dir", default="droit donnees", help="Source directory path.")
    parser.add_argument(
        "--output-dir", default="data/ingestion", help="Output directory for ingestion artifacts."
    )
    parser.add_argument(
        "--min-native-chars",
        type=int,
        default=120,
        help="If native PDF page text length is below this threshold, OCR fallback runs.",
    )
    parser.add_argument(
        "--ocr-dpi",
        type=int,
        default=200,
        help="Render DPI for OCR fallback (PDF page to image).",
    )
    parser.add_argument("--ocr-lang", default="fra", help="Tesseract OCR language code.")
    parser.add_argument(
        "--workers",
        type=int,
        default=DEFAULT_WORKERS,
        help=f"Number of parallel document workers (default: {DEFAULT_WORKERS}).",
    )
    parser.add_argument(
        "--tessdata-dir",
        default=None,
        help="Optional directory containing *.traineddata files (ex: data/ocr/tessdata).",
    )
    parser.add_argument(
        "--disable-ocr", action="store_true", help="Disable OCR fallback even when pages are low-text."
    )
    parser.add_argument(
        "--limit-docs",
        type=int,
        default=None,
        help="Optional limit for quick test runs.",
    )
    parser.add_argument(
        "--pdf-parser",
        choices=["auto", "native", "docling"],
        default="auto",
        help="PDF extraction backend: native parser/OCR, docling markdown parser, or auto fallback.",
    )

    ns = parser.parse_args()
    return IngestionArgs(
        input_dir=Path(ns.input_dir),
        output_dir=Path(ns.output_dir),
        min_native_chars=ns.min_native_chars,
        ocr_dpi=ns.ocr_dpi,
        ocr_lang=ns.ocr_lang,
        workers=max(1, int(ns.workers)),
        tessdata_dir=Path(ns.tessdata_dir) if ns.tessdata_dir else None,
        disable_ocr=ns.disable_ocr,
        limit_docs=ns.limit_docs,
        pdf_parser=str(ns.pdf_parser),
    )


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def to_jsonl(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def stable_doc_id(relative_path: str, sha256: str) -> str:
    seed = f"{relative_path}|{sha256}".encode("utf-8")
    return hashlib.sha1(seed).hexdigest()[:16]


def normalize_text(text: str) -> str:
    cleaned = text.replace("\x00", " ")
    lines = [line.strip() for line in cleaned.splitlines()]
    non_empty_lines = [line for line in lines if line]
    return "\n".join(non_empty_lines).strip()


def normalize_markdown(text: str) -> str:
    cleaned = text.replace("\x00", "")
    lines = [line.rstrip() for line in cleaned.splitlines()]
    merged = "\n".join(lines)
    merged = merged.replace("\r\n", "\n").replace("\r", "\n")
    merged = merged.strip()
    return merged


def resolve_tessdata_dir(user_tessdata_dir: Path | None) -> Path | None:
    if user_tessdata_dir and user_tessdata_dir.exists():
        return user_tessdata_dir

    env_dir = os.getenv("TESSDATA_PREFIX", "").strip()
    if env_dir and Path(env_dir).exists():
        return Path(env_dir)

    default_local = Path("data/ocr/tessdata")
    if default_local.exists():
        return default_local

    return None


def check_ocr_environment(
    ocr_lang: str,
    ocr_enabled: bool,
    user_tessdata_dir: Path | None,
) -> dict[str, Any]:
    result: dict[str, Any] = {
        "ocr_requested": ocr_enabled,
        "tesseract_in_path": False,
        "tesseract_path": None,
        "tessdata_dir": None,
        "tesseract_lang_available": False,
        "available_langs": [],
        "warnings": [],
    }
    if not ocr_enabled:
        return result

    tesseract_path = shutil.which("tesseract")
    if not tesseract_path and WINDOWS_TESSERACT_DEFAULT.exists():
        tesseract_path = str(WINDOWS_TESSERACT_DEFAULT)
    if not tesseract_path:
        result["warnings"].append("tesseract executable not found in PATH; OCR will be skipped.")
        return result

    result["tesseract_in_path"] = True
    result["tesseract_path"] = tesseract_path
    # Ensure pytesseract uses the detected executable, even if PATH is not set.
    pytesseract.pytesseract.tesseract_cmd = tesseract_path
    tessdata_dir = resolve_tessdata_dir(user_tessdata_dir)
    if tessdata_dir:
        result["tessdata_dir"] = str(tessdata_dir)

    try:
        env = os.environ.copy()
        if tessdata_dir:
            env["TESSDATA_PREFIX"] = str(tessdata_dir)
        proc = subprocess.run(
            [tesseract_path, "--list-langs"],
            capture_output=True,
            text=True,
            check=False,
            env=env,
        )
        lines = [line.strip() for line in proc.stdout.splitlines() if line.strip()]
        available = []
        for line in lines:
            if line.lower().startswith("list of available languages"):
                continue
            available.append(line)
        result["available_langs"] = sorted(set(available))
        result["tesseract_lang_available"] = ocr_lang in result["available_langs"]
        if not result["tesseract_lang_available"]:
            result["warnings"].append(
                f"language '{ocr_lang}' not found in tesseract list; OCR may fail."
            )
    except Exception as exc:  # noqa: BLE001
        result["warnings"].append(f"unable to query tesseract languages: {exc}")

    return result


def check_docling_environment(requested_parser: str) -> dict[str, Any]:
    installed = DocumentConverter is not None
    enabled = requested_parser == "docling" or (requested_parser == "auto" and installed)
    warnings: list[str] = []
    if requested_parser == "docling" and not installed:
        warnings.append("docling parser requested but package is not installed; fallback to native parser.")
    if requested_parser == "auto" and not installed:
        warnings.append("docling package not installed; using native parser.")
    return {
        "requested_parser": requested_parser,
        "docling_installed": installed,
        "docling_enabled": enabled and installed,
        "warnings": warnings,
    }


def build_manifest(args: IngestionArgs) -> list[dict[str, Any]]:
    if not args.input_dir.exists():
        raise FileNotFoundError(f"Input directory not found: {args.input_dir}")

    files = [
        path
        for path in args.input_dir.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    files.sort(key=lambda p: str(p).lower())
    if args.limit_docs is not None:
        files = files[: args.limit_docs]

    manifest: list[dict[str, Any]] = []
    for path in files:
        stat = path.stat()
        relative_path = path.relative_to(args.input_dir).as_posix()
        domain = relative_path.split("/", 1)[0] if "/" in relative_path else "root"
        file_sha256 = sha256_file(path)
        doc_id = stable_doc_id(relative_path, file_sha256)

        manifest.append(
            {
                "doc_id": doc_id,
                "relative_path": relative_path,
                "source_path": str(path),
                "domain": domain,
                "extension": path.suffix.lower(),
                "file_size_bytes": stat.st_size,
                "modified_utc": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
                "file_sha256": file_sha256,
            }
        )

    return manifest


def render_pdf_page_for_ocr(pdf: pdfium.PdfDocument, page_idx: int, dpi: int) -> Image.Image:
    page = pdf[page_idx]
    scale = dpi / 72.0
    bitmap = page.render(scale=scale)
    return bitmap.to_pil()


def render_fitz_page_for_ocr(doc: Any, page_idx: int, dpi: int) -> Image.Image:
    page = doc.load_page(page_idx)
    scale = dpi / 72.0
    matrix = fitz.Matrix(scale, scale)
    pix = page.get_pixmap(matrix=matrix, alpha=False)
    return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)


def ocr_image(img: Image.Image, ocr_lang: str, tessdata_dir: str | None) -> str:
    config_parts = ["--oem 3", "--psm 6"]
    if tessdata_dir:
        config_parts.insert(0, f"--tessdata-dir {Path(tessdata_dir).as_posix()}")
    config = " ".join(config_parts)
    try:
        text = pytesseract.image_to_string(img, lang=ocr_lang, config=config)
    finally:
        img.close()
    return normalize_text(text)


def ocr_pdf_page_with_pdfium(
    pdf: pdfium.PdfDocument,
    page_idx: int,
    ocr_lang: str,
    dpi: int,
    tessdata_dir: str | None,
) -> str:
    img = render_pdf_page_for_ocr(pdf, page_idx, dpi)
    return ocr_image(img, ocr_lang=ocr_lang, tessdata_dir=tessdata_dir)


def ocr_pdf_page_with_fitz(
    doc: Any,
    page_idx: int,
    ocr_lang: str,
    dpi: int,
    tessdata_dir: str | None,
) -> str:
    img = render_fitz_page_for_ocr(doc, page_idx, dpi)
    return ocr_image(img, ocr_lang=ocr_lang, tessdata_dir=tessdata_dir)


def extract_docx_pages(path: Path, doc_meta: dict[str, Any]) -> list[dict[str, Any]]:
    doc = Document(str(path))
    blocks = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    text = normalize_text("\n".join(blocks))

    return [
        {
            "doc_id": doc_meta["doc_id"],
            "relative_path": doc_meta["relative_path"],
            "source_path": doc_meta["source_path"],
            "domain": doc_meta["domain"],
            "extension": ".docx",
            "page_number": 1,
            "native_char_count": len(text),
            "ocr_char_count": 0,
            "ocr_used": False,
            "extraction_method": "native",
            "text": text,
        }
    ]


def extract_pdf_pages(
    path: Path,
    doc_meta: dict[str, Any],
    min_native_chars: int,
    ocr_enabled: bool,
    ocr_lang: str,
    ocr_dpi: int,
    tessdata_dir: str | None,
) -> tuple[list[dict[str, Any]], dict[str, int]]:
    reader = PdfReader(str(path))
    page_count = len(reader.pages)

    pdf_for_ocr: pdfium.PdfDocument | None = None
    fitz_for_ocr = None
    pdfium_open_error: str | None = None
    fitz_open_error: str | None = None

    if ocr_enabled:
        try:
            pdf_for_ocr = pdfium.PdfDocument(str(path))
        except Exception as exc:  # noqa: BLE001
            pdfium_open_error = str(exc)

    pages: list[dict[str, Any]] = []
    stats = {
        "page_count": page_count,
        "low_text_pages": 0,
        "ocr_attempted_pages": 0,
        "ocr_used_pages": 0,
        "ocr_error_pages": 0,
        "empty_pages_after_processing": 0,
    }

    for idx, page in enumerate(reader.pages):
        native_text = normalize_text(page.extract_text() or "")
        native_chars = len(native_text)
        final_text = native_text
        ocr_text = ""
        ocr_used = False
        method = "native"
        ocr_error = None

        should_try_ocr = ocr_enabled and native_chars < min_native_chars
        if native_chars < min_native_chars:
            stats["low_text_pages"] += 1

        if should_try_ocr:
            stats["ocr_attempted_pages"] += 1

            # Primary OCR renderer: PDFium (fast)
            if pdf_for_ocr is not None:
                try:
                    ocr_text = ocr_pdf_page_with_pdfium(
                        pdf_for_ocr,
                        idx,
                        ocr_lang=ocr_lang,
                        dpi=ocr_dpi,
                        tessdata_dir=tessdata_dir,
                    )
                except Exception as exc:  # noqa: BLE001
                    ocr_error = f"pdfium: {exc}"

            # Fallback renderer: PyMuPDF (more tolerant on malformed PDFs)
            if (not ocr_text) and fitz is not None:
                if fitz_for_ocr is None:
                    try:
                        fitz_for_ocr = fitz.open(str(path))
                    except Exception as exc:  # noqa: BLE001
                        fitz_open_error = str(exc)
                if fitz_for_ocr is not None:
                    try:
                        ocr_text = ocr_pdf_page_with_fitz(
                            fitz_for_ocr,
                            idx,
                            ocr_lang=ocr_lang,
                            dpi=ocr_dpi,
                            tessdata_dir=tessdata_dir,
                        )
                    except Exception as exc:  # noqa: BLE001
                        prefix = f"{ocr_error} | " if ocr_error else ""
                        ocr_error = f"{prefix}fitz: {exc}"

            if not ocr_text and ocr_error is None:
                open_errors = []
                if pdfium_open_error:
                    open_errors.append(f"pdfium_open: {pdfium_open_error}")
                if fitz_open_error:
                    open_errors.append(f"fitz_open: {fitz_open_error}")
                if not open_errors:
                    open_errors.append("OCR returned empty text")
                ocr_error = " | ".join(open_errors)

            if ocr_error:
                stats["ocr_error_pages"] += 1

            if len(ocr_text) > len(final_text):
                final_text = ocr_text
                ocr_used = True
                method = "ocr"
                stats["ocr_used_pages"] += 1

        if not final_text:
            stats["empty_pages_after_processing"] += 1

        pages.append(
            {
                "doc_id": doc_meta["doc_id"],
                "relative_path": doc_meta["relative_path"],
                "source_path": doc_meta["source_path"],
                "domain": doc_meta["domain"],
                "extension": ".pdf",
                "page_number": idx + 1,
                "native_char_count": native_chars,
                "ocr_char_count": len(ocr_text),
                "ocr_used": ocr_used,
                "ocr_error": ocr_error,
                "extraction_method": method,
                "text": final_text,
            }
        )

    if pdf_for_ocr is not None:
        pdf_for_ocr.close()
    if fitz_for_ocr is not None:
        fitz_for_ocr.close()

    return pages, stats


def extract_pdf_pages_docling(
    path: Path,
    doc_meta: dict[str, Any],
) -> tuple[list[dict[str, Any]], dict[str, int]]:
    if DocumentConverter is None:
        raise RuntimeError("docling is not installed")

    converter = DocumentConverter()
    result = converter.convert(str(path))
    document = getattr(result, "document", None)
    if document is None:
        raise RuntimeError("docling conversion returned no document object")

    markdown_text = ""
    if hasattr(document, "export_to_markdown"):
        markdown_text = str(document.export_to_markdown() or "")
    elif hasattr(document, "export_to_md"):
        markdown_text = str(document.export_to_md() or "")
    else:
        markdown_text = str(document)
    markdown_text = normalize_markdown(markdown_text)

    pages: list[dict[str, Any]] = []
    stats = {
        "page_count": 1,
        "low_text_pages": 0,
        "ocr_attempted_pages": 0,
        "ocr_used_pages": 0,
        "ocr_error_pages": 0,
        "empty_pages_after_processing": 0,
    }

    if not markdown_text:
        stats["empty_pages_after_processing"] = 1

    pages.append(
        {
            "doc_id": doc_meta["doc_id"],
            "relative_path": doc_meta["relative_path"],
            "source_path": doc_meta["source_path"],
            "domain": doc_meta["domain"],
            "extension": ".pdf",
            "page_number": 1,
            "native_char_count": len(markdown_text),
            "ocr_char_count": 0,
            "ocr_used": False,
            "ocr_error": None,
            "extraction_method": "docling_markdown",
            "text": markdown_text,
        }
    )
    return pages, stats


def process_single_document(
    doc: dict[str, Any],
    args: IngestionArgs,
    ocr_enabled: bool,
    tessdata_dir: str | None,
    docling_enabled: bool,
) -> tuple[list[dict[str, Any]], dict[str, Any], dict[str, int]]:
    source_path = Path(doc["source_path"])
    page_rows: list[dict[str, Any]] = []
    stats = {
        "page_count": 0,
        "low_text_pages": 0,
        "ocr_attempted_pages": 0,
        "ocr_used_pages": 0,
        "ocr_error_pages": 0,
        "empty_pages_after_processing": 0,
    }

    doc_report: dict[str, Any] = {
        "doc_id": doc["doc_id"],
        "relative_path": doc["relative_path"],
        "status": "ok",
        "error": None,
        "pages": 0,
        "low_text_pages": 0,
        "ocr_attempted_pages": 0,
        "ocr_used_pages": 0,
        "ocr_error_pages": 0,
        "empty_pages_after_processing": 0,
    }

    try:
        if doc["extension"] == ".pdf":
            use_docling = docling_enabled and args.pdf_parser in {"auto", "docling"}
            if use_docling:
                try:
                    page_rows, stats = extract_pdf_pages_docling(
                        source_path,
                        doc_meta=doc,
                    )
                except Exception as exc:  # noqa: BLE001
                    if args.pdf_parser == "docling":
                        raise
                    doc_report["status"] = "warning"
                    doc_report["error"] = f"docling failed, fallback native: {exc}"
                    page_rows, stats = extract_pdf_pages(
                        source_path,
                        doc_meta=doc,
                        min_native_chars=args.min_native_chars,
                        ocr_enabled=ocr_enabled,
                        ocr_lang=args.ocr_lang,
                        ocr_dpi=args.ocr_dpi,
                        tessdata_dir=tessdata_dir,
                    )
            else:
                page_rows, stats = extract_pdf_pages(
                    source_path,
                    doc_meta=doc,
                    min_native_chars=args.min_native_chars,
                    ocr_enabled=ocr_enabled,
                    ocr_lang=args.ocr_lang,
                    ocr_dpi=args.ocr_dpi,
                    tessdata_dir=tessdata_dir,
                )
        elif doc["extension"] == ".docx":
            page_rows = extract_docx_pages(source_path, doc_meta=doc)
            stats["page_count"] = len(page_rows)
        else:
            doc_report["status"] = "skipped"
    except Exception as exc:  # noqa: BLE001
        doc_report["status"] = "error"
        doc_report["error"] = str(exc)

    doc_report["pages"] = stats["page_count"] if stats["page_count"] else len(page_rows)
    doc_report["low_text_pages"] = stats["low_text_pages"]
    doc_report["ocr_attempted_pages"] = stats["ocr_attempted_pages"]
    doc_report["ocr_used_pages"] = stats["ocr_used_pages"]
    doc_report["ocr_error_pages"] = stats["ocr_error_pages"]
    doc_report["empty_pages_after_processing"] = stats["empty_pages_after_processing"]
    return page_rows, doc_report, stats


def ingest(args: IngestionArgs) -> dict[str, Any]:
    started_at = utc_now_iso()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    ocr_readiness = check_ocr_environment(
        args.ocr_lang,
        ocr_enabled=not args.disable_ocr,
        user_tessdata_dir=args.tessdata_dir,
    )
    ocr_enabled = (
        (not args.disable_ocr)
        and ocr_readiness["tesseract_in_path"]
        and ocr_readiness["tesseract_lang_available"]
    )
    tessdata_dir = ocr_readiness.get("tessdata_dir")
    docling_readiness = check_docling_environment(args.pdf_parser)
    docling_enabled = bool(docling_readiness["docling_enabled"])

    manifest = build_manifest(args)
    manifest_path = args.output_dir / "manifest.jsonl"
    to_jsonl(manifest_path, manifest)

    pages_out_path = args.output_dir / "pages.jsonl"
    doc_reports: list[dict[str, Any]] = []
    counters = {
        "docs_total": len(manifest),
        "pdf_docs": 0,
        "docx_docs": 0,
        "pages_total": 0,
        "low_text_pages": 0,
        "ocr_attempted_pages": 0,
        "ocr_used_pages": 0,
        "ocr_error_pages": 0,
        "empty_pages_after_processing": 0,
        "doc_errors": 0,
    }
    ordered_results: list[tuple[list[dict[str, Any]], dict[str, Any], dict[str, int]] | None] = [
        None
    ] * len(manifest)
    progress = tqdm(total=len(manifest), desc="Ingestion", unit="doc")

    if args.workers > 1:
        with ThreadPoolExecutor(max_workers=args.workers) as executor:
            futures = {
                executor.submit(
                    process_single_document,
                    doc,
                    args,
                    ocr_enabled,
                    tessdata_dir,
                    docling_enabled,
                ): idx
                for idx, doc in enumerate(manifest)
            }
            for future in as_completed(futures):
                idx = futures[future]
                ordered_results[idx] = future.result()
                progress.update(1)
    else:
        for idx, doc in enumerate(manifest):
            ordered_results[idx] = process_single_document(
                doc,
                args,
                ocr_enabled,
                tessdata_dir,
                docling_enabled,
            )
            progress.update(1)
    progress.close()

    with pages_out_path.open("w", encoding="utf-8") as pages_file:
        for idx, doc in enumerate(manifest):
            result = ordered_results[idx]
            if result is None:
                continue
            page_rows, doc_report, stats = result

            if doc["extension"] == ".pdf":
                counters["pdf_docs"] += 1
            elif doc["extension"] == ".docx":
                counters["docx_docs"] += 1

            for row in page_rows:
                pages_file.write(json.dumps(row, ensure_ascii=False) + "\n")

            counters["pages_total"] += doc_report["pages"]
            counters["low_text_pages"] += stats["low_text_pages"]
            counters["ocr_attempted_pages"] += stats["ocr_attempted_pages"]
            counters["ocr_used_pages"] += stats["ocr_used_pages"]
            counters["ocr_error_pages"] += stats["ocr_error_pages"]
            counters["empty_pages_after_processing"] += stats["empty_pages_after_processing"]
            if doc_report["status"] == "error":
                counters["doc_errors"] += 1

            doc_reports.append(doc_report)

    doc_report_path = args.output_dir / "doc_report.jsonl"
    to_jsonl(doc_report_path, doc_reports)

    report = {
        "started_at_utc": started_at,
        "finished_at_utc": utc_now_iso(),
        "input_dir": str(args.input_dir),
        "output_dir": str(args.output_dir),
        "config": {
            "min_native_chars": args.min_native_chars,
            "ocr_dpi": args.ocr_dpi,
            "ocr_lang": args.ocr_lang,
            "workers": args.workers,
            "ocr_requested": not args.disable_ocr,
            "ocr_enabled": ocr_enabled,
            "pdf_parser": args.pdf_parser,
            "tessdata_dir": str(args.tessdata_dir) if args.tessdata_dir else None,
            "limit_docs": args.limit_docs,
        },
        "ocr_readiness": ocr_readiness,
        "docling_readiness": docling_readiness,
        "stats": counters,
        "artifacts": {
            "manifest": str(manifest_path),
            "pages": str(pages_out_path),
            "doc_report": str(doc_report_path),
        },
    }

    report_path = args.output_dir / "ingestion_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main() -> None:
    args = parse_args()
    try:
        report = ingest(args)
    except Exception as exc:  # noqa: BLE001
        print(f"[ingestion] failed: {exc}", file=sys.stderr)
        raise

    stats = report["stats"]
    print("[ingestion] completed")
    print(f"docs_total={stats['docs_total']} pages_total={stats['pages_total']}")
    print(
        "low_text_pages="
        f"{stats['low_text_pages']} ocr_attempted_pages={stats['ocr_attempted_pages']} "
        f"ocr_used_pages={stats['ocr_used_pages']} ocr_error_pages={stats['ocr_error_pages']}"
    )
    print(
        "pdf_parser="
        f"{report['config']['pdf_parser']} docling_enabled={report['docling_readiness']['docling_enabled']}"
    )
    print(f"artifacts={report['artifacts']}")


if __name__ == "__main__":
    main()
